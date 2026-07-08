#!/usr/bin/env python3
"""
Generate academic paper Word document following Elsevier journal standards.
A4 paper, proper margins, Times New Roman / 宋体, 1.5 line spacing,
numbered sections, formatted tables with borders, justified text.

Usage:
    python3 tools/paper/generate_paper.py research/PAPER_CHINESE.md research/PAPER_CHINESE.docx \
        --lang zh \
        --title "FruitTreeScanner: 基于移动设备 LiDAR 的多模态果树果实检测与产量估算系统" \
        --authors "作者姓名" \
        --affiliations "单位信息" \
        --email "email@example.com" \
        --abstract "摘要内容..." \
        --keywords "果实检测；产量估算；LiDAR；点云；多模态融合"
"""

import re
import sys
import argparse
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ══════════════════════════════════════════════════════════════════════
#  Document Setup
# ══════════════════════════════════════════════════════════════════════

def setup_document(doc, language='en'):
    """Configure A4 page with academic margins and default fonts."""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)

    # Normal style — 12pt, justified, 1.5 spacing
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # East-Asian font support
    _set_rPr_font(style.element, ascii_font='Times New Roman',
                  east_asian_font='宋体', hansi_font='Times New Roman')

    # Heading styles — override ALL default heading colors to BLACK
    for level, (size, bold) in {1: (14, True), 2: (12, True),
                                  3: (11, True), 4: (11, False)}.items():
        hs = doc.styles[f'Heading {level}']
        _setup_heading_style(hs, size, bold, language)
        # Force black color on the style's font directly
        hs.font.color.rgb = RGBColor(0, 0, 0)


def _setup_heading_style(style, size, bold, language):
    font = style.font
    font.name = '黑体' if language == 'zh' else 'Times New Roman'
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = RGBColor(0, 0, 0)

    pf = style.paragraph_format
    pf.space_before = {14: Pt(18), 12: Pt(12), 11: Pt(8)}[size]
    pf.space_after = {14: Pt(8), 12: Pt(6), 11: Pt(4)}[size]
    pf.keep_with_next = True
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Force w:val="000000" at the XML level to override Word's theme color
    rPr = style.element.get_or_add_rPr()
    # Remove existing rFonts and color elements
    for child in list(rPr):
        if 'color' in child.tag or 'rFonts' in child.tag:
            rPr.remove(child)
    # Add our font and color
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} '
        f'w:ascii="Times New Roman" '
        f'w:eastAsia="黑体" '
        f'w:hAnsi="Times New Roman"/>'
    )
    rPr.append(rFonts)
    color_elm = parse_xml(
        f'<w:color {nsdecls("w")} w:val="000000"/>'
    )
    rPr.append(color_elm)


def _set_rPr_font(element, ascii_font, east_asian_font, hansi_font):
    rPr = element.get_or_add_rPr()
    # Remove existing rFonts if any
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} '
        f'w:ascii="{ascii_font}" '
        f'w:eastAsia="{east_asian_font}" '
        f'w:hAnsi="{hansi_font}"/>'
    )
    rPr.append(rFonts)


# ══════════════════════════════════════════════════════════════════════
#  Title Page
# ══════════════════════════════════════════════════════════════════════

def add_title_page(doc, title, authors, affiliations, email,
                   abstract_text, keywords_text, language='en'):
    """Insert title, authors, affiliations, abstract, keywords, then page break."""

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = '黑体' if language == 'zh' else 'Times New Roman'

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(authors)
    run.italic = True
    run.font.size = Pt(12)

    # Affiliations
    if affiliations:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(affiliations)
        run.font.size = Pt(10)
        run.italic = True

    # Email
    if email:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(24)
        run = p.add_run(email)
        run.font.size = Pt(10)

    # Abstract heading
    heading_label = '摘要' if language == 'zh' else 'Abstract'
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(heading_label)
    run.bold = True
    run.font.size = Pt(12)

    # Abstract body (single paragraph, no indent)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.first_line_indent = None
    run = p.add_run(abstract_text)
    run.italic = True
    run.font.size = Pt(11)

    # Keywords
    kw_label = '关键词' if language == 'zh' else 'Keywords'
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(f'{kw_label}：')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(keywords_text)
    run.font.size = Pt(11)

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════
#  Content Helpers
# ══════════════════════════════════════════════════════════════════════

def _add_formatted_text_run(paragraph, text, base_size=Pt(12)):
    """Append text to paragraph respecting **bold** inline markup."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.size = base_size


def _add_table(doc, headers, rows, caption=None):
    """Insert academic-style bordered table."""
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # Header
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shade_cell(cell, 'D9E2F3')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    # Data rows
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i + 1, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(cell_text)
            run.font.size = Pt(10)


def _shade_cell(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_disclaimer(doc, text):
    """Grey-block disclaimer paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)

    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _add_code_block(doc, code_lines):
    """Grey-block code listing."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)

    run = p.add_run('\n'.join(code_lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)


def _add_heading(doc, text, level):
    """Add section heading — always black text."""
    style_name = f'Heading {level}'
    p = doc.add_paragraph()
    p.style = doc.styles[style_name]
    run = p.add_run(text)
    run.bold = True
    run.font.size = {1: Pt(14), 2: Pt(12), 3: Pt(11), 4: Pt(11)}[level]
    run.font.color.rgb = RGBColor(0, 0, 0)  # Force black on the run


# ══════════════════════════════════════════════════════════════════════
#  Markdown → Docx Parser
# ══════════════════════════════════════════════════════════════════════

def parse_markdown_to_docx(md_file, docx_file, language='en',
                            title=None, authors=None, affiliations=None,
                            email=None, abstract_text=None, keywords_text=None):
    doc = Document()
    setup_document(doc, language)

    # ── Title page ────────────────────────────────────────────────
    if title and abstract_text:
        add_title_page(doc, title, authors or '', affiliations or '',
                       email or '', abstract_text, keywords_text or '',
                       language)

    # ── Read & parse body ─────────────────────────────────────────
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code = False
    code_buf = []
    table_rows = []

    for raw_line in lines:
        line = raw_line.rstrip('\n')

        # Code fence
        if line.startswith('```'):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                _add_code_block(doc, code_buf)
                code_buf = []
            continue

        if in_code:
            code_buf.append(line)
            continue

        # Table rows
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            if not all(c.startswith('---') for c in cells):
                table_rows.append(cells)
            continue
        elif table_rows:
            if len(table_rows) >= 2:
                _add_table(doc, table_rows[0], table_rows[1:])
                doc.add_paragraph('')
            table_rows = []
            continue

        # Blank line
        if not line.strip():
            doc.add_paragraph('')
            continue

        # Headings
        if line.startswith('# ') and not line.startswith('## '):
            _add_heading(doc, line[2:].strip(), 1)
            continue
        if line.startswith('## ') and not line.startswith('### '):
            _add_heading(doc, line[3:].strip(), 2)
            continue
        if line.startswith('### ') and not line.startswith('#### '):
            _add_heading(doc, line[4:].strip(), 3)
            continue
        if line.startswith('#### '):
            _add_heading(doc, line[5:].strip(), 4)
            continue

        # Blockquote → disclaimer
        if line.startswith('> '):
            _add_disclaimer(doc, line[2:].strip())
            continue

        # Standalone bold line
        stripped = line.strip()
        if stripped.startswith('**') and stripped.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip('* '))
            run.bold = True
            run.font.size = Pt(12)
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            p = doc.add_paragraph()
            p.style = doc.styles['List Number']
            _add_formatted_text_run(p, m.group(2).strip())
            continue

        # Bullet list
        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph()
            p.style = doc.styles['List Bullet']
            _add_formatted_text_run(p, line[2:].strip())
            continue

        # Regular paragraph
        if stripped:
            p = doc.add_paragraph()
            _add_formatted_text_run(p, stripped)

    doc.save(docx_file)
    print(f"✅  Saved → {docx_file}")


# ══════════════════════════════════════════════════════════════════════
#  CLI
# ══════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description='Generate academic paper Word (.docx) from Markdown')
    parser.add_argument('input', help='Input .md file')
    parser.add_argument('output', help='Output .docx file')
    parser.add_argument('--lang', choices=['en', 'zh'], default='en')
    parser.add_argument('--title', default='')
    parser.add_argument('--authors', default='')
    parser.add_argument('--affiliations', default='')
    parser.add_argument('--email', default='')
    parser.add_argument('--abstract', default='')
    parser.add_argument('--keywords', default='')
    args = parser.parse_args()

    parse_markdown_to_docx(
        args.input, args.output, args.lang,
        title=args.title, authors=args.authors,
        affiliations=args.affiliations, email=args.email,
        abstract_text=args.abstract, keywords_text=args.keywords,
    )


if __name__ == '__main__':
    main()
