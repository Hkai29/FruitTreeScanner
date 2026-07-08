#!/usr/bin/env python3
"""Convert Markdown to Word (.docx) with proper academic formatting"""

import re
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def format_heading(doc, level, text):
    """Format heading properly"""
    heading = doc.add_heading(text, level=level)
    
    # Configure fonts for heading
    for run in heading.runs:
        run.font.color.rgb = None  # Use default black
        run.font.size = {1: Pt(16), 2: Pt(14), 3: Pt(12), 4: Pt(11)}[level]
        if level <= 3:
            run.font.bold = True
        # Set Chinese font support
        rPr = run._element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:eastAsia="宋体" w:hAnsi="Times New Roman"/>')
        rPr.append(rFonts)
    
    # Configure spacing
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(12)
    elif level == 2:
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(10)
    elif level == 3:
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    
    return heading

def add_paragraph(doc, text, bold=False, italic=False, size=None, alignment=None, is_disclaimer=False):
    """Add formatted paragraph"""
    p = doc.add_paragraph()
    
    if is_disclaimer:
        # Disclaimer block with gray border background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
        p.paragraph_format.element.get_or_add_pPr().append(shading)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.right_indent = Cm(1)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
    
    # Process inline formatting: **bold** and *italic*
    # Split by ** first for bold
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            inner = part[2:-2]
            run = p.add_run(inner)
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic text
            inner = part[1:-1]
            run = p.add_run(inner)
            run.italic = True
        else:
            run = p.add_run(part)
        
        if bold and not run.bold:
            run.bold = bold
        if italic and not run.italic:
            run.italic = italic
        if size:
            run.font.size = size
        else:
            run.font.size = Pt(12) if not is_disclaimer else Pt(11)
        
        # Set font family
        rPr = run._element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:eastAsia="宋体" w:hAnsi="Times New Roman"/>')
        rPr.append(rFonts)
    
    if alignment:
        p.alignment = alignment
    elif is_disclaimer:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    return p

def add_table(doc, rows):
    """Add formatted table"""
    if not rows or len(rows) == 0:
        return
    
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table style
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            
            # Header row
            if i == 0:
                set_cell_shading(cell, "4472C4")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cell_text)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(11)
                rPr = run._element.get_or_add_rPr()
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:eastAsia="宋体" w:hAnsi="Times New Roman"/>')
                rPr.append(rFonts)
            else:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(cell_text)
                run.font.size = Pt(11)
                rPr = run._element.get_or_add_rPr()
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:eastAsia="宋体" w:hAnsi="Times New Roman"/>')
                rPr.append(rFonts)

def add_code_block(doc, code_lines, language=''):
    """Add code block with formatting"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    # Gray background
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    
    run = p.add_run('\n'.join(code_lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Courier New" w:hAnsi="Courier New"/>')
    rPr.append(rFonts)

def add_list_item(doc, text, level=0, numbered=False):
    """Add list item"""
    style_name = 'List Number' if numbered else 'List Bullet'
    p = doc.add_paragraph(text, style=style_name)
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    
    for run in p.runs:
        run.font.size = Pt(12)
        rPr = run._element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:eastAsia="宋体" w:hAnsi="Times New Roman"/>')
        rPr.append(rFonts)
    
    # Process inline formatting in list
    if '**' in text:
        p.clear()
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.size = Pt(12)
            rPr = run._element.get_or_add_rPr()
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:eastAsia="宋体" w:hAnsi="Times New Roman"/>')
            rPr.append(rFonts)
    
    return p

def convert_md_to_docx(md_file, docx_file):
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Configure paragraph spacing
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    
    # Set Chinese font support for normal style
    rPr = style.element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:eastAsia="宋体" w:hAnsi="Times New Roman"/>')
    rPr.append(rFonts)
    
    # Set margins (academic standard: 1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    in_code_block = False
    code_block_lines = []
    code_language = ''
    in_table = False
    table_rows = []
    
    for line in lines:
        # Code blocks
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_language = line[3:].strip()
                code_block_lines = []
            else:
                in_code_block = False
                add_code_block(doc, code_block_lines, code_language)
                code_block_lines = []
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            continue
        
        # Tables
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            # Skip separator row
            if all(c.startswith('---') or c.startswith('---') for c in cells if c):
                continue
            table_rows.append(cells)
            continue
        elif table_rows:
            add_table(doc, table_rows)
            table_rows = []
            doc.add_paragraph('')  # Spacing after table
            continue
        
        # Empty line
        if not line.strip():
            doc.add_paragraph('')
            continue
        
        # Headers
        if line.startswith('# ') and not line.startswith('##'):
            text = line[2:].strip()
            format_heading(doc, 1, text)
            continue
        elif line.startswith('## ') and not line.startswith('###'):
            text = line[3:].strip()
            format_heading(doc, 2, text)
            continue
        elif line.startswith('### ') and not line.startswith('####'):
            text = line[4:].strip()
            format_heading(doc, 3, text)
            continue
        elif line.startswith('#### '):
            text = line[5:].strip()
            format_heading(doc, 4, text)
            continue
        
        # Blockquotes (disclaimers)
        if line.startswith('> '):
            text = line[2:].strip()
            add_paragraph(doc, text, italic=True, is_disclaimer=True)
            continue
        
        # Bold line
        if line.startswith('**') and line.rstrip().endswith('**'):
            text = line.strip('* ')
            add_paragraph(doc, text, bold=True)
            continue
        
        # Numbered list
        numbered_match = re.match(r'^(\d+)\.\s+(.*)', line)
        if numbered_match:
            text = numbered_match.group(2).strip()
            add_list_item(doc, text, numbered=True)
            continue
        
        # Bullet list
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            add_list_item(doc, text)
            continue
        
        # Regular paragraph
        if line.strip():
            add_paragraph(doc, line.strip())
    
    doc.save(docx_file)
    print(f"✅ Saved: {docx_file}")

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python3 md2docx.py input.md output.docx")
        sys.exit(1)
    convert_md_to_docx(sys.argv[1], sys.argv[2])
